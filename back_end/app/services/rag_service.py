import os
import logging
from typing import List, Optional
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from typing import List as ListType
from langchain_google_genai import ChatGoogleGenerativeAI

from app.core.config import settings

logger = logging.getLogger(__name__)


# Removed RAGState as we're not using LangGraph anymore


class RAGService:
    def __init__(self, google_api_key: str):
        self.google_api_key = google_api_key
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-001",
            google_api_key=google_api_key
        )
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=google_api_key
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, 
            chunk_overlap=200
        )
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", "Bạn là một trợ lý AI chuyên trả lời câu hỏi dựa trên ngữ cảnh đã cho. "
                      "Hãy trả lời một cách chính xác và chi tiết dựa trên thông tin được cung cấp."),
            ("human", "Ngữ cảnh:\n{context}\n\nCâu hỏi: {question}")
        ])
        
    def create_vector_store(self, file_id: str) -> Chroma:
        """Create or get vector store for a specific file"""
        persist_directory = os.path.join(settings.storage_dir, "rag_data", file_id)
        os.makedirs(persist_directory, exist_ok=True)
        
        return Chroma(
            collection_name=f"doc_{file_id}",
            embedding_function=self.embeddings,
            persist_directory=persist_directory,
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text content from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error extracting text from TXT {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from supported file types"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def process_document(self, file_path: str, file_id: str) -> None:
        """Process document and store in vector database"""
        logger.info(f"Processing document for RAG: {file_path}")
        
        # Extract text content
        text_content = self.extract_text_from_file(file_path)
        
        if not text_content.strip():
            raise ValueError("Document contains no readable text content")
        
        # Create document object
        doc = Document(
            page_content=text_content,
            metadata={"source": file_path, "file_id": file_id}
        )
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        logger.info(f"Document split into {len(chunks)} chunks")
        
        # Store in vector database
        vector_store = self.create_vector_store(file_id)
        vector_store.add_documents(documents=chunks)
        
        logger.info(f"Document processed and stored in vector database")
    
    def retrieve_documents(self, file_id: str, question: str) -> List[Document]:
        """Retrieve relevant documents for a question"""
        try:
            logger.info(f"Creating vector store for file {file_id}")
            vector_store = self.create_vector_store(file_id)
            logger.info(f"Performing similarity search for: {question}")
            retrieved_docs = vector_store.similarity_search(question, k=5)
            logger.info(f"Found {len(retrieved_docs)} similar documents")
            return retrieved_docs
        except Exception as e:
            logger.error(f"Error in retrieve_documents: {e}")
            raise
    
    def generate_answer(self, question: str, context_docs: List[Document]) -> str:
        """Generate answer based on retrieved context"""
        docs_content = "\n\n".join(doc.page_content for doc in context_docs)
        messages = self.prompt.invoke({
            "question": question, 
            "context": docs_content
        })
        response = self.llm.invoke(messages)
        return response.content
    
    def query_document(self, file_id: str, question: str) -> str:
        """Query a processed document with a question"""
        try:
            import asyncio
            
            logger.info(f"Querying document {file_id} with question: {question}")
            
            # Handle event loop for FastAPI context
            def _run_rag():
                # Ensure we have an event loop in this thread
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    # No event loop, create one
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                
                # Simple RAG pipeline
                context_docs = self.retrieve_documents(file_id, question)
                logger.info(f"Retrieved {len(context_docs)} documents")
                
                answer = self.generate_answer(question, context_docs)
                logger.info(f"Generated answer: {answer[:100]}...")
                
                return answer
            
            return _run_rag()
                
        except Exception as e:
            logger.error(f"Error querying document {file_id}: {e}")
            raise


def get_rag_service(google_api_key: str) -> RAGService:
    """Get RAG service instance"""
    return RAGService(google_api_key)
